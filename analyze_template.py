#!/usr/bin/env python3
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

doc = Document(r'Izvestaj\conference-template-a4.docx')

print("=== TEMPLATE ANALIZA ===\n")
print(f"Broj paragrafa: {len(doc.paragraphs)}\n")

for i, p in enumerate(doc.paragraphs[:30]):
    txt = p.text[:70] if p.text else "[PRAZNO]"
    align = p.alignment
    bold_info = "BOLD" if p.runs and any(r.bold for r in p.runs) else "normal"
    size_info = [f"{r.font.size.pt}pt" for r in p.runs if r.font.size] if p.runs else []
    center = "CENTER" if align == WD_ALIGN_PARAGRAPH.CENTER else "LEFT" if align == WD_ALIGN_PARAGRAPH.LEFT else str(align)
    
    print(f"{i:2d}: {txt:<70} | {bold_info:6s} | {center:10s} | Size: {size_info}")

# Proveri sekcije
print("\n=== SEKCIJE ===")
for sec in doc.sections:
    cols = sec._sectPr.xpath('./w:cols')
    if cols:
        num_cols = cols[0].get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}num')
        print(f"Kolone: {num_cols}")
