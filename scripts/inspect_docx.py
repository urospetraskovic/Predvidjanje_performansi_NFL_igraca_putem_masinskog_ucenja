"""Inspect the existing Izvestaj docx to learn styles."""
from docx import Document
from pathlib import Path

SRC = Path(r"c:\Users\Korisnik\Documents\GitHub\Analiza-i-Obrada\Izvestaj\Predviđanje_performansi_NFL_igrača_putem_mašinskog_učenja_Milan_Jovkić_Uroš_Petrašković.docx")

d = Document(str(SRC))
print("Paragraphs:", len(d.paragraphs))
print("Sections:", len(d.sections))
print("Tables:", len(d.tables))

print("\n--- Used styles (paragraph) ---")
used_styles = {}
for p in d.paragraphs:
    used_styles[p.style.name] = used_styles.get(p.style.name, 0) + 1
for k, v in sorted(used_styles.items(), key=lambda x: -x[1]):
    print(f"  {k}: {v}")

print("\n--- First 80 paragraphs ---")
for i, p in enumerate(d.paragraphs[:80]):
    txt = p.text[:120].replace('\n', ' ')
    print(f"  [{i}] style={p.style.name!r:30s} | {txt}")

print("\n--- Last 30 paragraphs ---")
for i, p in enumerate(d.paragraphs[-30:]):
    txt = p.text[:120].replace('\n', ' ')
    print(f"  [-{30-i}] style={p.style.name!r:30s} | {txt}")

print("\n--- All heading-like paragraphs (all) ---")
for i, p in enumerate(d.paragraphs):
    if 'Heading' in p.style.name or 'Naslov' in p.style.name or p.style.name.startswith('Title'):
        print(f"  [{i}] style={p.style.name!r:30s} | {p.text[:150]}")
