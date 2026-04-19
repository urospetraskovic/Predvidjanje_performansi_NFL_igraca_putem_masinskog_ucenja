"""Replace the existing 'NEURONSKE MREŽE' chapter in the report with a richer
narrative version that embeds loss-curve figures and result tables.

Inputs:
- Source docx (transitional namespace) at REPORT_IN
- Figure folder at FIG_DIR

Output:
- Edited docx written to REPORT_OUT
"""
from __future__ import annotations
import copy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

REPORT_IN = Path(
    r"c:\Users\Korisnik\Documents\GitHub\Analiza-i-Obrada\Izvestaj\_report_transitional.docx"
)
REPORT_OUT = Path(
    r"c:\Users\Korisnik\Documents\GitHub\Analiza-i-Obrada\Izvestaj\_report_NN_v2.docx"
)
FIG_DIR = Path(r"c:\Users\Korisnik\Documents\GitHub\Analiza-i-Obrada\results")

# ------------------------------------------------------------------
# 1. open document and locate insertion anchors
# ------------------------------------------------------------------
doc = Document(str(REPORT_IN))
body = doc.element.body

children = list(body)


def find_paragraph(text: str):
    for el in children:
        if el.tag != qn("w:p"):
            continue
        t = "".join(t.text or "" for t in el.iter(qn("w:t")))
        if t.strip() == text:
            return el
    raise KeyError(f"Paragraph not found: {text!r}")


nn_heading = find_paragraph("NEURONSKE MREŽE")
literatura = find_paragraph("LITERATURA")

# remove every element between nn_heading (inclusive) and literatura (exclusive)
to_remove = []
in_section = False
for el in children:
    if el is nn_heading:
        in_section = True
    if in_section:
        if el is literatura:
            break
        to_remove.append(el)
for el in to_remove:
    body.remove(el)

# ------------------------------------------------------------------
# 2. helpers for inserting content before LITERATURA
# ------------------------------------------------------------------
anchor = literatura  # everything we add will go right before this element


def add_paragraph(text: str, style: str = "Body Text"):
    p = doc.add_paragraph(text, style=style)
    body.remove(p._p)
    anchor.addprevious(p._p)
    return p


def add_heading(text: str, level: int):
    style = f"Heading {level}"
    return add_paragraph(text, style=style)


def add_caption(text: str):
    return add_paragraph(text, style="Caption")


def add_image(filename: str, width_cm: float = 14.5, caption: str | None = None):
    src = FIG_DIR / filename
    if not src.exists():
        # fall back: just note it in the doc
        add_paragraph(f"[Slika nedostaje: {filename}]", style="Body Text")
        if caption:
            add_caption(caption)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(src), width=Cm(width_cm))
    body.remove(p._p)
    anchor.addprevious(p._p)
    if caption:
        add_caption(caption)


def add_table(rows):
    """rows: list[list[str]]; first row is header."""
    n_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i].cells[j]
            cell.text = ""
            par = cell.paragraphs[0]
            r = par.add_run(str(val))
            if i == 0:
                r.bold = True
                par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    body.remove(table._tbl)
    anchor.addprevious(table._tbl)
    return table


# ------------------------------------------------------------------
# 3. write the new NN chapter
# ------------------------------------------------------------------

# ---- H1 + intro -------------------------------------------------
add_heading("NEURONSKE MREŽE", 1)

add_paragraph(
    "Nakon detaljnog razmatranja klasičnih regresionih modela na sezonski "
    "agregiranim podacima, u ovom poglavlju rad se proširuje primenom neuronskih "
    "mreža. Neuronske mreže su parametarski modeli inspirisani biološkim neuronima "
    "koji uče nelinearne preslikavanje ulaznih obeležja u izlaz kompozicijom "
    "većeg broja jednostavnih nelinearnih transformacija (sloj × težinska matrica × "
    "aktivacija). Treniranje se sprovodi minimizacijom funkcije gubitka algoritmom "
    "stohastičkog gradijentnog spusta i njegovim varijantama (Adam, AdamW), uz "
    "propagaciju greške unazad (backpropagation). Za razliku od linearnih modela i "
    "tree ansambala, neuronske mreže nemaju ugrađen eksplicitni mehanizam za "
    "izbor informativnih obeležja niti za hvatanje vremenskih obrazaca, pa je "
    "neophodan pažljiv izbor arhitekture i sistematska regularizacija."
)

add_paragraph(
    "Najznačajnija promena u odnosu na prethodni deo rada jeste promena ciljne "
    "promenljive: umesto sezonski agregiranih jardi po utakmici (yards/G) sada se "
    "predviđaju ostvareni prijemni jardi (receiving_yards) Wide Receiver-a po "
    "pojedinačnoj utakmici. Razlog za ovaj prelazak je dvojak. Prvi je obim podataka "
    "— neuronske mreže imaju veliki broj parametara i stabilno generalizuju tek na "
    "skupovima reda veličine 10⁴–10⁵ uzoraka. Sezonski skup za WR poziciju sadrži "
    "samo 5.529 zapisa, što je premalo i za umereno duboke arhitekture; prelaskom "
    "na nivo utakmice raspolažemo sa 46.115 zapisa za 1.719 unikatnih igrača kroz "
    "11 sezona (2015–2025), što je više od osam puta veći skup. Drugi razlog je da "
    "ovakav target sa eksplicitnom vremenskom strukturom (niz utakmica po igraču) "
    "otvara prostor za primenu rekurentnih arhitektura (LSTM, GRU), koje su "
    "projektovane upravo za sekvencijalne podatke."
)

add_paragraph(
    "Pregled distribucije nove ciljne promenljive (mean = 29,6; std = 31,8; max = "
    "300; oko 12% utakmica sa 0 jardi) potvrđuje da je dovoljno bogata i dovoljno "
    "varijabilna za netrivijalan zadatak učenja. Istovremeno, izrazita desna "
    "asimetrija i prisustvo legitimnih boom utakmica nameću potrebu za pažljivim "
    "izborom funkcije gubitka i regularizacionih tehnika, jer obični MSE "
    "prekomerno kažnjava velike, ali realne, ishode."
)

add_paragraph(
    "Da bi se dobilo direktno poređenje, na istom game-level skupu obučeni su i "
    "klasični modeli mašinskog učenja iz prethodnog dela rada (RandomForest, "
    "XGBoost, LightGBM, ElasticNet), uz identičan protokol evaluacije. Kao "
    "neuronske arhitekture testirani su MLP modeli različitog kapaciteta i jačine "
    "regularizacije kao osnovni (base) feed-forward pristup, kao i dve klase "
    "rekurentnih mreža — GRU i LSTM — sa standardnim i proširenim varijantama. "
    "Konačno, testirani su ansambli koji kombinuju klasične i neuronske modele, "
    "kao i kvantilna regresija za predviđanje intervala poverenja."
)

# ---- H2: Metodologija i implementacija ----
add_heading("Metodologija i implementacija", 2)

add_paragraph(
    "Metodološki okvir zasniva se na četiri stuba: striktna vremenska podela "
    "podataka radi izbegavanja curenja informacija iz budućnosti, bogato "
    "inženjerstvo obeležja koje kodira kratkoročnu i dugoročnu formu igrača, skup "
    "arhitektura neuronskih mreža prilagođenih tabularnim i sekvencijalnim "
    "podacima i sistematska optimizacija hiperparametara pomoću Optune. Svi "
    "modeli dele isti protokol evaluacije kako bi rezultati bili neposredno uporedivi."
)

# H3: skup podataka
add_heading("Skup podataka i vremenska podela", 3)
add_paragraph(
    "Korišćen je skup wr_all_weeks.csv koji sadrži 46.115 zapisa na nivou "
    "igrač–utakmica za 1.719 Wide Receiver-a iz 11 sezona (2015–2025), sa 99 "
    "originalnih kolona. Ciljna promenljiva je receiving_yards, sa značajnom desnom "
    "asimetrijom (mean = 29,6; std = 31,8; max = 300) i sa oko 12% utakmica u kojima "
    "igrač ostvari nula jardi, što unosi dodatnu inflated-zero komponentu. U odnosu "
    "na sezonski skup iz prethodnog dela rada (5.529 redova za WR, target yards/G), "
    "novi skup je više od osam puta veći i ima eksplicitnu temporalnu strukturu na "
    "nivou pojedinačnih utakmica."
)
add_paragraph(
    "Podela je striktno hronološka, bez ikakvog mešanja: trening 2015–2021 (29.276 "
    "uzoraka), validacija 2022–2023 (8.921 uzorak) i test 2024–2025 (6.199 uzoraka). "
    "Test skup služi isključivo za konačnu evaluaciju, nikada za odabir arhitekture ili "
    "hiperparametara. Kao i u prethodnom delu rada, primenjen je StandardScaler "
    "obučen isključivo na trening skupu, uz ranije opisane procedure imputacije "
    "nedostajućih vrednosti i kodiranja kategoričkih kolona. Za ciljnu promenljivu "
    "primenjene su dve transformacije: √y (najbolje radi sa MLP-om) i log(1+y) "
    "(stabilnije za rekurentne arhitekture), oba izbora izabrana su empirijski na "
    "osnovu validacionog MAE-a."
)
add_paragraph(
    "Zbog izražene asimetrije targeta, uvedena je težinska shema uzoraka po "
    "formuli w = 1 + 0,6·√(y / mean_y), što daje opseg težina [1,00; 2,87]. Ovakvo "
    "ponderisanje sprečava model da konzistentno potcenjuje igrače koji povremeno "
    "ostvare boom utakmice, jer im se u funkciji gubitka dodeljuje proporcionalno "
    "veći značaj."
)

# H3: feature engineering
add_heading("Inženjerstvo obeležja i kreiranje sekvenci", 3)
add_paragraph(
    "Za razliku od stabala odlučivanja, neuronske mreže ne mogu samostalno "
    "ekstrahovati vremenske obrasce iz tabularnih kolona. Eksplicitno su zato "
    "konstruisana obeležja koja kodiraju različite vremenske horizonte forme "
    "igrača. Inženjering je sproveden u četiri sloja."
)
add_paragraph(
    "Bazne statistike: receiving_yards, targets, receptions, epa, target_share, "
    "air_yards_share i pregame informacije (spread, total, domaćin/gost, defanzivni "
    "rang protivnika). Temporalne transformacije: za svako bazno obeležje generisano "
    "je šest derivata — _lag1 (vrednost iz prethodne utakmice), _roll3 i _roll5 "
    "(rolling proseci poslednje 3 odnosno 5 utakmica), _career_avg (ekspandirajući "
    "karijerni prosek), _expanding_std (karijerna standardna devijacija kao mera "
    "volatilnosti) i _momentum (trend definisan kao roll5 − lag1). Sve transformacije "
    "koriste shift(1) kako tekuća utakmica ne bi ulazila u svoje sopstveno obeležje. "
    "Interakciona obeležja: target_share × pregame_total hvata sinergiju volumena i "
    "tempa, weeks_since_last_game kodira pauze i povrede, is_new_season je "
    "indikator prve utakmice nove sezone — važan za rekurentne mreže koje prelaze "
    "granice sezona. Matchup obeležja: defanzivna statistika protivnika izračunata "
    "kao ekspandirajući prosek iz prethodnih nedelja u sezoni; isključivo prošli "
    "podaci garantuju odsustvo curenja."
)
add_paragraph(
    "Ukupno je konstruisano 184 inženjerisana obeležja. Rani eksperimenti su "
    "pokazali da kompletan skup izaziva preprilagođavanje neuronskih modela zbog "
    "visoke korelisanosti rolling statistika, pa je sproveden korak selekcije "
    "obeležja kombinovanjem XGBoost i LightGBM ocena važnosti. Top 40 obeležja "
    "(51,8% kumulativne važnosti) sačuvano je u datoteci selected_features_top40."
    "json i korišćeno kao primarni ulaz svih kasnijih modela. Pet najznačajnijih "
    "obeležja su first_downs_roll5, target_share_std_lag1, targets_roll5, "
    "air_yards_roll5 i pregame_total."
)
add_image(
    "feature_importance_top40.png",
    caption="Slika 11. Top 40 obeležja po kombinovanoj XGBoost i LightGBM oceni važnosti.",
)
add_image(
    "feature_selection_overfit_gap.png",
    caption="Slika 12. Razlika između trening i test R² za pun (184 obeležja) i redukovan (40 obeležja) skup. Vidljivo je smanjenje overfit gap-a za neuronske modele nakon selekcije.",
)
add_paragraph(
    "Za rekurentne modele, pored tabularnih obeležja, formirane su i sekvence po "
    "igraču. Svaka sekvenca je hronološki niz utakmica jednog igrača; za svaki "
    "vremenski korak čuva se 33 sirova game-level obeležja (bez lag/roll "
    "transformacija, jer RNN sam uči temporalne obrasce), dok 20 statičkih pregame "
    "obeležja (spread, total, domaćin/gost, defanzivni rang) ulaze odvojeno kroz "
    "drugu granu modela. Da kratke karijere ne bi bile odbacivane, koristi se padding "
    "neutralnom vrednošću 0,0 i Masking sloj koji ignoriše padded korake. Maksimalna "
    "dužina sekvence (T) varirana je između 5 i 12 utakmica; finalni izbor T = 12 "
    "pokrio je medijan dužine karijere bez eksplozije udela padded koraka."
)

# H3: arhitekture
add_heading("Arhitekture modela", 3)
add_paragraph(
    "Testirano je nekoliko klasa neuronskih arhitektura, redom od najjednostavnije "
    "do najkompleksnije. Svaka klasa odgovara konkretnoj hipotezi o tome šta ograničava "
    "performansu prethodnih modela: MLP kao tabularna referenca, GRU/LSTM zbog "
    "sekvencijalne prirode podataka, kvantilna regresija zbog asimetrije i autlajera, "
    "arhitektura sa pažnjom i ugrađivanjem identiteta igrača kao najsloženiji pokušaj, "
    "i ansambli kao finalni korak konsolidacije."
)

add_heading("Feed-forward mreže (MLP) kao bazni model", 4)
add_paragraph(
    "Osnovni MLP (Model B u nastavku) prihvata 40 selektovanih obeležja kao ravan "
    "vektor i prosleđuje ih kroz pet potpuno povezanih slojeva dimenzija "
    "[448, 128, 320, 448, 256] sa ReLU aktivacijama, LayerNorm normalizacijom i "
    "Dropoutom 0,5 nakon svakog sloja. Izlazni sloj je jedan neuron sa linearnom "
    "aktivacijom koji predviđa √(receiving_yards). Bias inicijalizovan je nulama, a "
    "težine HeNormal raspodelom. Varijante Model C, D i E uvode dodatne "
    "regularizacione mehanizme: GaussianNoise na ulazu (σ ∈ [0,05; 0,25]) za "
    "stabilizaciju gradijenata, Mixup augmentaciju primera (α ∈ [0,1; 0,4]) i "
    "kompaktnije arhitekture (3–4 sloja) sa ciljem smanjenja preprilagođavanja bez "
    "gubitka kapaciteta."
)

add_heading("Rekurentne mreže (GRU i LSTM)", 4)
add_paragraph(
    "Rekurentne mreže su ključna klasa modela za podatke sa sekvencijalnom "
    "strukturom. Testirane su dve varijante ćelija: LSTM (Long Short-Term Memory) "
    "[22], koja koristi tri kapije (input, forget, output) i odvojeno ćelijsko stanje, "
    "i GRU (Gated Recurrent Unit) [23], jednostavnija arhitektura sa dve kapije "
    "(reset i update) i bez odvojenog ćelijskog stanja. Inicijalna postavka "
    "(WR_GRU_Baseline) koristi jedan GRU(32) sloj sa Dense(1) izlazom nad sliding "
    "window sekvencama dužine 5 i 95 sirovih obeležja po koraku — namerno "
    "minimalna konfiguracija kao referentna donja granica."
)
add_paragraph(
    "Ključni redizajn — Improved RNN — uveo je dvoulaznu (dual-input) arhitekturu: "
    "sekvencijalna grana prima T × 33 sirova game-level obeležja, prolazi kroz "
    "Masking(0,0), zatim jedan LSTM(128) sloj (unidirekcioni), pa LayerNorm i "
    "Dropout(0,30); statička grana paralelno prima 20 pregame obeležja, prolazi "
    "kroz GaussianNoise(0,20), Dense, LayerNorm i Dropout. Dva izlaza se konkatenizuju "
    "i prolaze kroz završni Dense → izlazni neuron. Ovakva podela omogućava "
    "rekurentnoj grani da uči striktno temporalni signal, dok se pregame informacije "
    "(koje su konstantne tokom utakmice) tretiraju kao statički kontekst. Identična "
    "konfiguracija testirana je i sa GRU ćelijama (LSTM vs GRU ablacija)."
)

add_heading("Pažnja i ugrađivanje identiteta igrača", 4)
add_paragraph(
    "Najsloženija testirana arhitektura ima tri ulazne grane: sekvencijalnu (T × 45 "
    "obeležja, gde je 45 = 33 sirova + 6 lag1 + 6 roll3 inženjerisanih kanala), "
    "embedding granu (player_id → Embedding(1308, 8) sa L2 regularizacijom 1e-4) i "
    "statičku granu (20 pregame obeležja). Sekvencijalna grana koristi Bidirectional "
    "LSTM/GRU sa return_sequences=True; izlaz prolazi kroz prilagođeni "
    "AttentionPool sloj koji uči težine za sve vremenske korake — umesto da koristi "
    "samo poslednje skriveno stanje, softmax preko naučenih skorova (uz maskiranje "
    "padded pozicija sa −∞) omogućava modelu da fokusira pažnju na najrelevantnije "
    "utakmice u istoriji igrača. Sve tri grane spajaju se konkatenacijom i ulaze u "
    "završni Dense blok (96 → 1)."
)

add_heading("Kvantilna regresija (MLP Quantile)", 4)
add_paragraph(
    "Zbog asimetrije i autlajera u distribuciji jardi, testiran je i pristup kvantilne "
    "regresije. Arhitektura koristi istu MLP-B osnovu ([448, 128, 320, 448, 256]), ali "
    "sa Dense(3) izlazom koji simultano predviđa kvantile q10, q50 i q90 ciljne "
    "promenljive. Trening koristi pinball (quantile) loss L(y, ŷ, τ) = max(τ·(y−ŷ), "
    "(τ−1)·(y−ŷ)). Kao tačkasta predikcija koristi se medijan q50, koji je po "
    "konstrukciji L1-optimalan i znatno robusniji na outliere od srednje vrednosti. "
    "Kao dodatna prednost, razmak q90 − q10 daje direktan interval poverenja od "
    "80%, koji se može koristiti i za fantasy/kladioničarsku primenu."
)

add_heading("Strategije ansambliranja", 4)
add_paragraph(
    "Pojedinačni modeli dostižu sličan kvalitet, ali njihove greške nisu savršeno "
    "korelisane (koeficijenti korelacije reziduala u opsegu 0,96–0,99), što ostavlja "
    "prostor za marginalni dobitak kroz ansamblovanje. Testirane su četiri strategije: "
    "simple_mean (prosta aritmetička sredina predikcija), inverse_mae (težine obrnuto "
    "proporcionalne validacionom MAE-u), constrained_ls (nenegativni najmanji "
    "kvadrati sa ograničenjem da je suma težina = 1) i ridge_stack (Ridge meta-učnik "
    "obučen nad predikcijama baznih modela). Testirana su dva sastava: Opcija 1 = "
    "{MLP Hybrid, MLP Quantile, LightGBM, ElasticNet} i Opcija 2 = {XGBoost, "
    "LightGBM, MLP Hybrid}."
)

# H3: trening
add_heading("Proces treniranja", 3)
add_paragraph(
    "Optimizator je AdamW sa weight decay-em u opsegu [1e-4; 5e-3]; klasičan Adam "
    "se u praksi pokazao dramatično lošiji jer L2 dodatak primenjuje na momente "
    "umesto na same težine. Kao funkcija gubitka dominantno se koristi Huber loss "
    "sa delta parametrom biranim Optunom u opsegu [0,5; 2,0] — Huber spaja "
    "kvadratnu osetljivost MSE-a oko nule sa linearnom robusnošću MAE-a dalje od "
    "nule, što odgovara distribuciji jardi gde su outlieri legitimni boom uzorci. Za "
    "kvantilnu regresiju koristi se pinball loss, a za najranije baseline modele (GRU "
    "baseline) klasičan MSE."
)
add_paragraph(
    "Regularizacioni paket uključuje Dropout (0,2–0,5 u zavisnosti od sloja), "
    "LayerNorm (forsiran u svim modernim arhitekturama; ako se prepusti Optuni, "
    "ona u polovini slučajeva bira 'none' i dobije lošiji rezultat), GaussianNoise na "
    "ulazu (σ = 0,05–0,25), Mixup augmentaciju trening batcha (α ∈ [0,1; 0,4]), weight "
    "decay i L2 regularizaciju na embedding sloju (1e-4). Batch size variraju od 32 "
    "do 256 zavisno od arhitekture. Tri ugrađena Keras callback-a su konstantno "
    "aktivna: ModelCheckpoint (monitor = val_loss, save_best_only = True), "
    "EarlyStopping (patience 12–30, restore_best_weights = True) i "
    "ReduceLROnPlateau (factor = 0,3; patience = 5–10; min_lr = 1e-5). Svi "
    "eksperimenti reproducibilni su zahvaljujući fiksiranim seed-ovima za python, "
    "numpy i tensorflow (seed = 42)."
)

# H3: optuna
add_heading("Optimizacija hiperparametara primenom Optuna biblioteke", 3)
add_paragraph(
    "Za razliku od klasičnih modela kod kojih je Grid Search bio dovoljan, neuronske "
    "mreže imaju znatno bogatiji prostor pretrage: dubina mreže, broj jedinica po "
    "sloju, dropout po sloju, learning rate, weight decay, batch size, delta za Huber, "
    "α za Mixup, σ za GaussianNoise, dužina sekvence za RNN, broj rekurentnih "
    "slojeva, opcija bidirektionalnosti, izbor LayerNorm-a. Primenjena je Bayesova "
    "optimizacija kroz Optuna biblioteku sa TPE (Tree-structured Parzen Estimator) "
    "samplerom i MedianPruner-om koji agresivno prekida loše probe nakon "
    "minimum tri epohe. Standardna konfiguracija je 100 trials po modelu, sa "
    "objective funkcijom min(val MAE), čime se direktno optimizuje primarna metrika "
    "interesa. Keras Tuner korišćen je u ranim iteracijama, ali je u kasnijim fazama u "
    "potpunosti zamenjen Optunom zbog fleksibilnijeg API-ja, podrške za "
    "early-stopping trials i mogućnosti definisanja proizvoljnih distribucija."
)
add_table(
    [
        ["Klasa modela", "Hiperparametar", "Opseg / Skup vrednosti"],
        ["MLP", "broj slojeva", "{2, 3, 4, 5}"],
        ["MLP", "broj jedinica po sloju", "{64, 128, 192, 256, 320, 384, 448}"],
        ["MLP", "dropout", "[0,2; 0,5]"],
        ["MLP", "learning_rate", "[1e-4; 1e-2] (log)"],
        ["MLP", "weight_decay", "[1e-4; 5e-3] (log)"],
        ["MLP", "Huber delta", "[0,5; 2,0]"],
        ["MLP Hybrid", "GaussianNoise σ", "[0,05; 0,25]"],
        ["MLP Hybrid", "Mixup α", "[0,1; 0,4]"],
        ["RNN", "tip ćelije", "{LSTM, GRU}"],
        ["RNN", "dužina sekvence T", "{5, 6, 8, 10, 12}"],
        ["RNN", "broj rekurentnih slojeva", "{1, 2}"],
        ["RNN", "broj jedinica", "{32, 64, 96, 128, 160, 192}"],
        ["RNN", "dropout / recurrent_dropout", "[0,1; 0,4]"],
        ["RNN", "ciljna transformacija", "{√y, log(1+y)}"],
        ["Quantile MLP", "kvantili (fiksno)", "{0,10, 0,50, 0,90}"],
        ["Sve klase", "batch size", "{32, 64, 128, 256}"],
        ["Sve klase", "broj epoha (max)", "{200, 300}"],
        ["Sve klase", "EarlyStopping patience", "{12, 20, 30}"],
    ]
)
add_caption("Tabela 7. Prostor pretrage hiperparametara po klasi arhitektura.")

# ---- H2: rezultati i diskusija ----
add_heading("Rezultati i diskusija", 2)

add_paragraph(
    "Razvoj neuronskih modela tekao je iterativno, kroz devet uzastopnih faza "
    "podeljenih u ukupno 22 Jupyter notebook-a. Svaka naredna faza motivisana je "
    "konkretnim nalazom prethodne (najčešće utvrđenim overfit-om ili lošom "
    "generalizacijom), što je omogućilo ciljano poboljšavanje slabih tačaka pre "
    "uvođenja kompleksnijih arhitektura. U nastavku je prikazana hronologija "
    "eksperimenata sa pratećim grafikonima loss funkcije, zatim sumarna tabela svih "
    "modela, pa direktno poređenje sa klasičnim referentnim modelima i konačno "
    "dijagnostička analiza preostale greške."
)

# H3: faza 1 - baseline + overfit
add_heading("Bazni RNN i pojava overfit-a", 3)
add_paragraph(
    "Najjednostavniji moguć RNN — jedan GRU sloj sa 32 jedinice, sekvenca dužine "
    "5 i 95 sirovih obeležja — služi kao referentna donja granica iznad koje svaka "
    "naredna arhitektura mora da se dokaže. Postiže R² = 0,2590 na test skupu, što "
    "je daleko ispod kvaliteta klasičnih modela. Već u ovoj fazi vidi se osnovni "
    "problem: trening loss veoma brzo opada, dok validacioni loss rano stagnira i "
    "počinje da raste, što je klasičan obrazac preprilagođavanja."
)
add_image(
    "gru_training_history.png",
    caption="Slika 13. Trening i validacioni loss baznog GRU modela kroz epohe.",
)
add_paragraph(
    "Paralelno je u objedinjenom pipeline-u sa Optunom (100 trials po modelu) "
    "testiran čitav set arhitektura — RandomForest, XGBoost, LightGBM, MLP, TCN i "
    "multi-layer RNN. Tree modeli postižu R² ≈ 0,33–0,335, MLP R² ≈ 0,316, dok "
    "rekurentne mreže ostaju daleko ispod (TCN 0,245; LSTM/GRU 0,216). Razlika R² ≈ "
    "0,12 između LightGBM-a i RNN-a je signal da je arhitektura RNN-a u ovoj postavci "
    "fundamentalno pogrešno dimenzionisana. Ovaj rezultat je pokretač svih daljih "
    "iteracija."
)
add_image(
    "lstm_training_history.png",
    caption="Slika 14. Trening i validacioni loss multi-layer LSTM modela iz Faze 2 — vidljiv jasan overfit gap već posle 10 epoha.",
)

# H3: feature reduction
add_heading("Smanjenje broja obeležja kao prvi alat protiv overfit-a", 3)
add_paragraph(
    "Prvi sistematski korak za smanjenje preprilagođavanja jeste agresivna selekcija "
    "obeležja. Kombinovana XGBoost + LightGBM analiza važnosti pokazala je da od "
    "184 inženjerisanih obeležja, 144 efektivno predstavljaju šum za neuronske "
    "modele — visoko su korelisana sa boljim varijantama (npr. roll3 i roll5 verzije "
    "iste statistike) ili imaju vrlo malu marginalnu informacionu vrednost. Redukcija "
    "na top 40 obeležja poboljšava R² MLP-a za +0,021 i istovremeno smanjuje overfit "
    "gap (razlika trening i validacionog R²) za 0,01–0,03 zavisno od modela. "
    "Klasični tree modeli ne dobijaju značajno na ovoj redukciji, što je konzistentno "
    "sa njihovim ugrađenim mehanizmom selekcije."
)
add_image(
    "feature_selection_learning_curves.png",
    caption="Slika 15. Trening i validacione krive MLP-a za pun (184) i redukovan (40) skup obeležja — vidljivo je sporije divergiranje krivih u redukovanoj varijanti.",
)
add_image(
    "feature_selection_comparison.png",
    caption="Slika 16. Poređenje test metrika (MAE i R²) baznih modela za pun i redukovan skup obeležja.",
)

# H3: MLP variants
add_heading("MLP varijante: trade-off kapaciteta i regularizacije", 3)
add_paragraph(
    "Pet kontrolisanih MLP varijanti (A–E) izoluje efekat broja obeležja i jačine "
    "regularizacije. Model A koristi sva 184 obeležja sa istom arhitekturom kao "
    "Model B (40 obeležja); razlika u R² (0,308 vs 0,329) potvrđuje pozitivan uticaj "
    "selekcije. Model C uvodi kompaktniju arhitekturu (3 sloja [256,128,64]) sa "
    "GaussianNoise(0,2), Mixup(α=0,3) i softer sample weights — postiže najmanji "
    "overfit gap, ali po cenu malo nižeg R²-a. Model D (Fixed Hybrid) zadržava punu "
    "arhitekturu Modela B i dodaje GaussianNoise(0,15) i Mixup(0,2); rezultat je "
    "najniži MAE (17,83) među svim kontrolnim MLP-ovima. Model E (Optuna Hybrid, "
    "100 trials) bira 4 sloja [320,384,192,128] sa noise=0,25 i mixup=0,2, dajući "
    "najbolji balans MAE/R² (17,88; 0,3225)."
)
add_image(
    "mlp_comparison_learning_curves.png",
    caption="Slika 17. Trening i validacione krive za MLP modele A, B i C.",
)
add_image(
    "mlp_comparison_overfit_gap.png",
    caption="Slika 18. Overfit gap (razlika trening i val R²) po MLP varijantama. Vidljivo je smanjenje gap-a od Modela A ka Modelu C.",
)
add_image(
    "mlp_hybrid_learning_curves.png",
    caption="Slika 19. Trening i validacione krive Hybrid varijanti (Model D i E) sa GaussianNoise i Mixup-om — krive su znatno bliže nego kod Modela A.",
)
add_paragraph(
    "Generalni nalaz iz ove serije: veći kapacitet bez regularizacije daje viši R² ali "
    "lošiji MAE; jača regularizacija (noise + mixup + manja arhitektura) daje niži "
    "MAE ali blago niži R². Ovaj trade-off ostaje konstanta u svim daljim fazama i "
    "kasnije se reflektuje i na izbor strategije ansamblovanja."
)

# H3: RNN redesign
add_heading("Redizajn rekurentnih mreža (Improved RNN)", 3)
add_paragraph(
    "Loš rezultat ranih RNN modela motivisao je radikalan redizajn arhitekture, koji "
    "je doveo do prvog rekurentnog modela koji se može meriti sa tree modelima. "
    "Promenjeno je šest stvari istovremeno. Prvo, sekvencijalna grana sada prima 33 "
    "sirova game-level obeležja po koraku umesto 184 lag/roll inženjerisanih, jer "
    "RNN sam uči temporalne obrasce iz raw signala. Drugo, 20 pregame obeležja "
    "ulazi odvojeno kroz statičku Dense granu, što znači da mreža ne mora svako "
    "konstantno obeležje da pamti unutar skrivenog stanja. Treće, kratke karijere se "
    "više ne odbacuju (ranije je seq_len ≥ T filter izbacivao 7–17 hiljada uzoraka); "
    "umesto toga koristi se padding nulama i Masking sloj. Četvrto, arhitektura je "
    "namerno smanjena — najviše dva LSTM sloja sa najviše 192 jedinice — jer "
    "Optuna na velikom search space-u tipično bira preveliki model koji overfit-uje. "
    "Peto, LayerNorm je forsiran (ne prepušten Optuni); šesto, GaussianNoise je "
    "premešten isključivo na statičku granu (jer šum nad padded pozicijama u "
    "sekvenciji razbija detekciju maske)."
)
add_image(
    "rnn_improved_learning_curves.png",
    caption="Slika 20. Trening i validacione krive Improved RNN modela — pad val loss-a ide gotovo paralelno sa trening loss-om i nastavlja da opada do epohe 35–40.",
)
add_image(
    "rnn_improved_full_comparison.png",
    caption="Slika 21. Test metrike Improved RNN-a u poređenju sa baznim GRU/LSTM modelima i sa Optuna-tuniranim multi-layer RNN-om iz Faze 2.",
)
add_paragraph(
    "Rezultat je dramatičan: MAE = 18,03; RMSE = 25,11; R² = 0,3354 — Improved "
    "RNN postaje prvi neuronski model koji blago nadmašuje LightGBM (R² = 0,3346). "
    "U odnosu na inicijalni multi-layer RNN, R² je porastao za 0,1197, MAE pao za "
    "1,70 jardi. Zaključak: za ovaj problem, jednostavnija arhitektura sa "
    "sačuvanim podacima i pravilnom maskom dramatično nadmašuje veliki Optuna-"
    "biran model koji baca kratke karijere i koristi šum na sekvenciji."
)
add_paragraph(
    "Direktna LSTM vs GRU ablacija (isti hiperparametri, izmenjena samo ćelija) "
    "pokazuje da je razlika minimalna: LSTM postiže R² = 0,3354, GRU @ LSTM-HP "
    "0,3344, GRU Optuna-best 0,3326. Razlika u R² je manja od 0,003, što znači da "
    "tip rekurentne ćelije nije ključan — ključna je sama topologija (dual-input + "
    "padding/masking + forsiran LayerNorm)."
)
add_image(
    "rnn_gru_vs_lstm_comparison.png",
    caption="Slika 22. LSTM vs GRU ablacija — razlika u test metrikama je u granicama merne neizvesnosti.",
)

# H3: attention experiment
add_heading("Pažnja i ugrađivanje identiteta igrača: limit pristupa", 3)
add_paragraph(
    "Sledeća iteracija pokušala je da proširi Improved RNN dodavanjem dva napredna "
    "mehanizma — pažnje (AttentionPool) preko cele istorije i ugrađivanja identiteta "
    "igrača (Embedding(1308, 8)). Hipoteza je bila da self-attention može da "
    "pondera različite utakmice u istoriji, a embedding da nauči latentne "
    "kvalitete pojedinih igrača (npr. vrhunski WR vs prosečan WR). Testirana je grid "
    "{LSTM, GRU} × {√y, log(1+y)} = 4 modela."
)
add_image(
    "rnn_attn_embed_learning_curves.png",
    caption="Slika 23. Trening i validacione krive Attention + Embedding modela. Svi modeli rano staju (epoha 3–4), što ukazuje na overfit ili preagresivnu regularizaciju.",
)
add_paragraph(
    "Rezultati su ispod Improved RNN-a (najbolji GRU_sqrt R² = 0,3236). Diagnostika "
    "(WR_RNN_Attention_PlayerEmbed_v2 sa relaxiranom regularizacijom — dropout "
    "0,25 → 0,15, noise 0,05 → 0, weight decay 5e-4 → 1e-4, patience 25 → 50) "
    "blago popravlja stvar (GRU_log1p_v2 R² = 0,3262), ali svi modeli i dalje staju "
    "rano. Uzrok je dvostruk: 49,9% test uzoraka ima player_id koji se ne pojavljuje "
    "u trening skupu (Out-of-Vocabulary problem), pa je embedding sloj efektivno "
    "beskoristan na test skupu; takođe, vrlo mali broj utakmica po igraču (medijan "
    "≈ 14) ne pruža dovoljno signala da self-attention nauči smislene temporalne "
    "težine."
)
add_image(
    "rnn_attn_embed_v2_learning_curves.png",
    caption="Slika 24. Krive nakon relaxirane regularizacije u Attention + Embedding v2 — dijagnostika potvrđuje da je arhitektura, ne regularizacija, glavno ograničenje.",
)
add_paragraph(
    "Zaključak ove faze: za naš dataset, jednostavniji Improved RNN ostaje bolji "
    "izbor; pažnja i embedding ne donose očekivano poboljšanje, jer ograničenje nije "
    "u kapacitetu modela već u količini i strukturi raspoloživih podataka."
)

# H3: quantile
add_heading("Kvantilna regresija — najniži MAE među pojedinačnim modelima", 3)
add_paragraph(
    "Pristup kvantilne regresije pokazao se kao najuspešnija pojedinačna tehnika za "
    "primarnu metriku (MAE). Ista MLP-B osnova obučava se sa pinball loss-om za "
    "kvantile q10, q50 i q90. Tačkasta predikcija je medijan q50, koji je po "
    "konstrukciji L1-optimalan i znatno robusniji od srednje vrednosti. Pored "
    "predikcije, model proizvodi i koristan dodatak: q90 − q10 raspon kao 80% "
    "interval poverenja."
)
add_image(
    "mlp_quantile_learning_curves.png",
    caption="Slika 25. Trening i validacione krive MLP Quantile modela — sva tri kvantila konvergiraju stabilno, bez vidljivog overfit-a.",
)
add_image(
    "mlp_quantile_predictions.png",
    caption="Slika 26. Predikcije medijane (q50) i pojas q10–q90 na test skupu — pokrivenost je 78,8% (cilj 80%), širina pojasa 59,7 jardi.",
)
add_paragraph(
    "Rezultati: MAE q50 = 17,76; RMSE = 25,60; R² = 0,3093; pokrivenost q10–q90 = "
    "78,8% (cilj 80%); prosečna širina pojasa 59,7 jardi. MLP Quantile je najbolji "
    "pojedinačan model po MAE metrici i predstavlja primer kako pažljivo izabrana "
    "funkcija gubitka može dati bolji rezultat od arhitekturnih komplikacija."
)

# H3: ensemble
add_heading("Strategije ansambliranja", 3)
add_paragraph(
    "Pet baznih modela (XGBoost, LightGBM, ElasticNet, MLP Hybrid, MLP Quantile) "
    "kombinovano je u četiri ansambl strategije. Pre kombinovanja, izračunata je "
    "matrica korelacije njihovih reziduala — vrednosti u opsegu 0,964–0,993 "
    "fundamentalno ograničavaju potencijalnu dobit jer modeli prave slične greške."
)
add_image(
    "ensemble_final_residual_corr.png",
    caption="Slika 27. Matrica korelacije reziduala pet baznih modela. Visoka međusobna korelacija ograničava maksimalnu dobit ansambla.",
)
add_paragraph(
    "Ipak, dobitak postoji. Najbolja strategija po MAE-u je constrained_ls "
    "(nenegativni najmanji kvadrati sa sumom težina = 1) sa sastavom Opcija 1, koja "
    "dodeljuje 70,4% težine MLP Quantile-u i 19,2% LightGBM-u. Rezultat: MAE = "
    "17,708 — najniži ostvaren u celokupnom istraživanju. Najbolji R² = 0,3384 "
    "ostaje stari ansambl (XGB + LGB + MLP) na 184 obeležja iz Faze 2, što je još "
    "jedan dokaz trade-off-a između MAE i R² (constrained_ls žrtvuje deo R² u "
    "korist ekstremno niskog MAE)."
)
add_image(
    "ensemble_final_comparison.png",
    caption="Slika 28. Poređenje četiri ansambl strategije sa najboljim pojedinačnim modelima — simple_mean i constrained_ls dele najbolju ukupnu poziciju.",
)

# H3: sumarni rezultati
add_heading("Sumarni rezultati svih modela", 3)
add_paragraph(
    "Tabela u nastavku sumarizuje ključne neuronske i ansambl modele na test skupu "
    "(2024–2025, 6.199 uzoraka) sortirano po MAE — primarnoj metrici optimizacije."
)
add_table(
    [
        ["Model", "Obeležja", "Test MAE", "Test RMSE", "Test R²"],
        ["Ensemble Opt1 / constrained_ls", "40", "17,708", "25,430", "0,3184"],
        ["MLP Quantile q50", "40", "17,75", "25,60", "0,3093"],
        ["Ensemble Opt1 / simple_mean", "40", "17,798", "25,243", "0,3284"],
        ["MLP Hybrid (Model D)", "40", "17,83", "25,67", "0,3055"],
        ["RandomForest (game-level)", "40", "17,85", "—", "0,3321"],
        ["MLP Hybrid Optuna (Model E)", "40", "17,88", "25,35", "0,3225"],
        ["MLP Improved (Model C)", "40", "17,91", "25,49", "0,3151"],
        ["LSTM Improved (najbolji RNN)", "33+20", "18,03", "25,11", "0,3354"],
        ["Stari ansambl (XGB+LGB+MLP)", "184", "18,05", "25,05", "0,3384"],
        ["XGBoost", "40", "18,09", "25,498", "0,3148"],
        ["MLP Original (Model A)", "184", "18,10", "25,62", "0,3079"],
        ["LightGBM", "40", "18,11", "25,659", "0,3061"],
        ["GRU Improved", "33+20", "18,14", "25,13", "0,3344"],
        ["MLP Optuna (multi-layer)", "184", "18,16", "25,48", "0,3156"],
        ["GRU log1p + Attention v2", "45+20", "18,29", "25,28", "0,3262"],
        ["TCN (Dilated Conv1D)", "184", "19,30", "27,59", "0,2453"],
        ["Multi-layer RNN (Faza 2)", "184", "19,73", "27,93", "0,2157"],
        ["Bazni GRU", "95", "20,31", "28,54", "0,2590"],
    ]
)
add_caption(
    "Tabela 8. Rezultati neuronskih modela i ansambla na test skupu (receiving_yards, originalna skala)."
)
add_paragraph(
    "Najbolji rezultati po metrikama: najniži MAE = 17,708 ostvaruje ansambl Opt1 / "
    "constrained_ls; najviši R² = 0,3384 stari ansambl XGB+LGB+MLP iz Faze 2; "
    "najbolji pojedinačni MAE 17,75 MLP Quantile q50; najbolji rekurentni model R² "
    "= 0,3354 LSTM Improved. Devet od deset najboljih modela koristi tačno 40 "
    "selektovanih obeležja, što potvrđuje ključni uticaj selekcije obeležja na "
    "generalizaciju neuronskih mreža."
)

# H3: comparison with baseline
add_heading("Poređenje sa klasičnim referentnim modelima", 3)
add_paragraph(
    "Direktno numeričko poređenje sa modelima iz prethodnog dela rada nije moguće "
    "jer prethodni rezultati (RandomForest R² ≈ 0,325 za WR poziciju u single-output "
    "postavci) se odnose na sezonski target yards/G, čija je varijansa drastično "
    "manja od game-level varijanse. Zato su isti klasični modeli (RandomForest, "
    "XGBoost, LightGBM, ElasticNet) ponovo obučeni na game-level zadatku pod "
    "identičnim protokolom kao i neuronske mreže, što omogućava korektno "
    "poređenje."
)
add_table(
    [
        ["Model", "Klasa", "Test MAE", "Test R²"],
        ["RandomForest", "klasični", "17,85", "0,3321"],
        ["XGBoost", "klasični", "18,09", "0,3148"],
        ["LightGBM", "klasični", "18,11", "0,3061"],
        ["ElasticNet", "klasični", "18,10", "0,3192"],
        ["MLP Quantile q50", "neuronski", "17,75", "0,3093"],
        ["MLP Hybrid (Model D)", "neuronski", "17,83", "0,3055"],
        ["LSTM Improved", "neuronski", "18,03", "0,3354"],
        ["Ensemble Opt1 / constrained_ls", "ansambl", "17,708", "0,3184"],
        ["Ensemble (XGB+LGB+MLP, Faza 2)", "ansambl", "18,05", "0,3384"],
    ]
)
add_caption(
    "Tabela 9. Poređenje neuronskih i klasičnih modela na istom game-level zadatku."
)
add_paragraph(
    "Tri zaključka. Prvo, pojedinačne neuronske mreže (MLP Quantile, MLP Hybrid, "
    "LSTM Improved) u najboljem slučaju dostižu ili blago nadmašuju najbolje "
    "klasične modele po MAE-u, ali razlika je marginalna. Drugo, tree modeli "
    "postižu viši R² od najboljeg pojedinačnog neuronskog modela, što sugeriše da "
    "su nelinearnosti koje hvata stablo dobro razdvojene od onih koje hvata duboka "
    "mreža. Treće, ansambli koji kombinuju obe klase modela donose konzistentno "
    "najbolje rezultate, što potvrđuje hipotezu o različitim tipovima grešaka. Iz "
    "metodološkog ugla, na ovom konkretnom problemu neuronske mreže ne donose "
    "kvantno preimućstvo nad tree modelima, jer tabularni podaci sa visokim šumom "
    "u targetu nemaju bogatu strukturu kao slika ili tekst."
)
add_image(
    "error_metrics_comparison.png",
    caption="Slika 29. Poređenje MAE i R² po klasama modela na game-level zadatku.",
)

# H3: error by bin
add_heading("Analiza grešaka po segmentima targeta", 3)
add_paragraph(
    "Da bi se utvrdilo gde tačno modeli greše, predikcije su analizirane u četiri "
    "opsega stvarnog broja jardi (0–30, 30–60, 60–100, 100+). Tabela poredi "
    "najbolji pojedinačni model (MLP Quantile) i najbolji ansambl (Opt1 / "
    "constrained_ls)."
)
add_table(
    [
        ["Opseg (yards)", "N", "MAE — MLP Quantile", "MAE — Ensemble", "Δ (Ens − MLP)"],
        ["0–30", "3.837", "11,24", "11,57", "−0,33"],
        ["30–60", "1.359", "18,21", "17,70", "+0,51"],
        ["60–100", "677", "33,96", "33,06", "+0,90"],
        ["100–300", "239", "73,74", "72,53", "+1,21"],
    ]
)
add_caption(
    "Tabela 10. MAE po opsezima stvarnog broja jardi. Pozitivno Δ znači da je ansambl bolji."
)
add_paragraph(
    "Greške su izrazito heterogene. Za niske ishode (0–30 jardi, 61,9% test "
    "skupa) MLP Quantile je bolji od ansambla jer njegov medijan efikasno pokriva "
    "dominantni mod distribucije. Za srednje i visoke ishode (>30 jardi) ansambl "
    "preuzima prednost, sa rastućom dobit (+0,51, +0,90, +1,21), što znači da "
    "kombinovanje modela donosi najveću vrednost upravo za boom utakmice. "
    "Apsolutna greška na intervalu 100+ jardi je više puta veća od greške na niskim "
    "ishodima, što reflektuje i teorijski plafon (single-game šum) i mali broj "
    "uzoraka u tom opsegu (239 od 6.199)."
)

# H3: ceiling
add_heading("Teorijski plafon i ceiling analiza", 3)
add_paragraph(
    "Centralno pitanje istraživanja je da li R² ≈ 0,33 predstavlja plafon naših "
    "modela ili plafon samog problema. Postojanje takvog plafona testirano je serijom "
    "ceiling eksperimenata u kojima je originalni single-game target zamenjen "
    "centered rolling mean targetom sa prozorima 2, 3, 5 i 7 utakmica. Hipoteza je "
    "jednostavna — ako modeli na izglađenom targetu postižu drastično bolji R², "
    "znači da postoji značajan single-game šum koji se ne može predvideti i da "
    "originalni rezultati dostižu plafon problema, ne plafon modela."
)
add_table(
    [
        ["Prozor (W)", "Test MAE (smooth)", "Test R² (smooth)", "Test MAE (raw)", "Test R² (raw)"],
        ["sirov target", "—", "—", "18,03", "0,3354"],
        ["2-game", "12,10", "0,5840", "18,48", "0,2993"],
        ["3-game", "11,12", "0,6018", "18,68", "0,3114"],
        ["5-game", "8,85", "0,7119", "19,17", "0,2988"],
        ["7-game", "7,73", "0,7626", "19,41", "0,2911"],
    ]
)
add_caption(
    "Tabela 11. Ceiling analiza: performanse LSTM Improved na centered rolling mean targetu različitog prozora (W=2, 3, 5, 7) prema izglađenom i prema sirovom targetu."
)
add_image(
    "wr_ceiling_combined_r2.png",
    caption="Slika 30. R² po veličini smoothing prozora — vidljiv monoton porast od 0,33 (sirov) do 0,76 (7-game).",
)
add_image(
    "wr_ceiling_combined.png",
    caption="Slika 31. MAE po veličini smoothing prozora — pad sa 18 jardi na 7,7 jardi pri 7-game prozoru.",
)
add_paragraph(
    "Zaključak ceiling analize je nedvosmislen. Kada se single-game šum ukloni "
    "usrednjavanjem, isti modeli postižu R² do 0,76 i MAE od svega 7,7 jardi na "
    "7-game prozoru. Evaluacija sa glatkim predikcijama prema sirovom (negladenom) "
    "targetu ostaje praktično ista kao i originalni rezultati (R² 0,29–0,34, MAE "
    "18,5–19,4) — što znači da modeli ekstrahuju gotovo sav signal koji uopšte "
    "postoji za single-game predikciju. Dalje poboljšanje zahteva kvalitativno nove "
    "informacije (real-time injury status, snap count projekcije, inside matchup "
    "data), a ne arhitekturna doterivanja. Dodatna verifikacija sa per-split "
    "smoothing-om (notebook WR_5_game_Ceiling_Smoothed_NoLeakage) potvrdila je "
    "da nema skrivenog cross-split curenja informacija — razlika u rezultatima manja "
    "je od 0,01 R²."
)

# H3: feature ablation
add_heading("Ablacija obeležja: opadajući povraćaji", 3)
add_paragraph(
    "Dodatna dijagnostička analiza sproverena je kroz treniranje finalnih modela "
    "samo sa najvažnijih 5 i 10 obeležja, kako bi se kvantifikovalo koliki deo signala "
    "potiče iz nekoliko najinformativnijih kolona."
)
add_table(
    [
        ["Skup obeležja", "MLP Quantile MAE", "Δ vs Top-40"],
        ["Top-5", "17,87", "+0,12"],
        ["Top-10", "17,78", "+0,03"],
        ["Top-40", "17,75", "—"],
    ]
)
add_caption("Tabela 12. Ablacija broja obeležja — krivulja opadajućih povraćaja.")
add_paragraph(
    "Samo top-5 obeležja nosi veliku većinu signala — MAE se pogoršava za svega "
    "0,12 jarda u odnosu na top-40, a prelaz sa top-10 na top-40 donosi dodatnih "
    "samo 0,03 jarda. Ovo jasno demonstrira opadajuće povraćaje i dodatno "
    "potvrđuje da je odabir obeležja pomoću ansambla tree modela metodološki "
    "ispravan i da je top-40 skup pragmatičan kompromis između informativnosti i "
    "kompaktnosti."
)

# ---- H2: zakljucak ----
add_heading("Zaključak", 2)
add_paragraph(
    "Primena dubokog učenja na game-level predviđanje prijemnih jardi NFL Wide "
    "Receiver-a donela je nekoliko ključnih nalaza. Prelazak sa sezonskih na "
    "game-level podatke obezbedio je više od osam puta veći skup (5.529 → 46.115 "
    "uzoraka), čime je opravdana primena neuronskih arhitektura. Bazni RNN modeli "
    "su odmah pokazali jasan overfit, koji je sistematski tretiran u tri koraka — "
    "selekcijom 40 najznačajnijih obeležja umesto 184, smanjenjem arhitekture "
    "(maks. 2 rekurentna sloja sa najviše 192 jedinice) i pažljivim regularizacionim "
    "paketom (LayerNorm, Dropout, weight decay, GaussianNoise samo na statičkoj "
    "grani, Mixup samo bez padding-a). Najveću pojedinačnu razliku napravio je "
    "redizajn RNN-a u dual-input arhitekturu (sirov sekvencijalni signal + odvojen "
    "statički kontekst, sa padding + masking-om), koji je R² rekurentnih mreža "
    "podigao sa 0,2157 na 0,3354 — dovoljno da se izjednače sa najboljim tree "
    "modelima."
)
add_paragraph(
    "Kompleksnije arhitekture sa pažnjom i ugrađivanjem identiteta igrača nisu "
    "donele dodatno poboljšanje, jer ograničenje nije u kapacitetu modela već u "
    "količini podataka po igraču (medijan ≈ 14 utakmica) i u 49,9% Out-of-Vocabulary "
    "stope na test skupu. Kvantilna regresija (MLP Quantile) izdvojila se kao "
    "najuspešnija pojedinačna tehnika — dala je najniži MAE (17,75) i besplatne "
    "intervale poverenja sa 78,8% pokrivenošću, što je primenljivo i za fantasy i za "
    "kladioničarski kontekst. Strategije ansambliranja (constrained_ls i simple_mean) "
    "dodatno su smanjile MAE do 17,708 — apsolutni rekord — iako visoka "
    "korelisanost reziduala (0,96–0,99) između baznih modela ograničava maksimalnu "
    "dobit ansambla."
)
add_paragraph(
    "Sistematska analiza otkrila je jasan trade-off između MAE i R²: modeli sa "
    "većim kapacitetom i manjom regularizacijom imaju viši R² (Model B, stari "
    "ansambl), dok jače regularizovani ili L1-optimalni modeli (MLP Quantile, "
    "constrained_ls ansambl) postižu niži MAE. Preporuka za praksu zavisi od "
    "primene: za rangiranje igrača prioritet je R² (simple_mean ansambl ili stari "
    "XGB + LGB + MLP), dok je za tačnu numeričku procenu jarda prioritet MAE "
    "(constrained_ls ansambl ili MLP Quantile samostalno)."
)
add_paragraph(
    "Najvažniji opšti nalaz dolazi iz ceiling analize. Kada se single-game šum "
    "ukloni usrednjavanjem (centered rolling mean prozori 2, 3, 5 i 7 utakmica), isti "
    "modeli postižu R² do 0,76 i MAE od 7,7 jardi, dok evaluacija prema sirovom "
    "targetu ostaje na 18–19 jardi. Drugim rečima, R² u opsegu 0,33–0,37 nije "
    "plafon naših modela, već fundamentalna granica single-game predikcije sa "
    "raspoloživim obeležjima. Modeli ekstrahuju gotovo sav dostupan signal, a "
    "preostali ireducibilni šum (standardna devijacija single-game greške ≈ 17–23 "
    "jardi) potiče iz inherentne varijabilnosti NFL utakmica — povreda u toku "
    "utakmice, dinamičkog game script-a, taktičkih izmena protivničke odbrane — "
    "koje se ne mogu predvideti čak ni iz pažljivo inženjerisanih pregame "
    "obeležja. Ovo je istovremeno i najveće dostignuće rada (jasna kvantifikacija "
    "plafona) i najveće ograničenje (taj plafon je nizak)."
)
add_paragraph(
    "Metodološki, ovo proširenje pokazuje koliko je važna disciplina koja kombinuje "
    "inženjerstvo obeležja, selekciju obeležja, redizajn arhitekture vođen "
    "dijagnostikom (a ne slepim dodavanjem kompleksnosti), sistematsku Bayesovu "
    "optimizaciju hiperparametara i dijagnostičku proveru granica problema. Ovaj niz "
    "koraka, primenjen uz striktnu vremensku podelu i fiksirane seed-ove, "
    "predstavlja prenosiv šablon za druge sportske analitičke probleme sa visokim "
    "nivoom šuma. Konačno, kombinacija klasičnih i neuronskih modela u okviru "
    "ansambla Opt1 / constrained_ls daje najmanju srednju apsolutnu grešku od "
    "17,71 jardi po utakmici na 6.199 testnih uzoraka, što predstavlja kvantitativni "
    "optimum ostvariv u okviru opisane metodologije."
)

# ------------------------------------------------------------------
# 4. save
# ------------------------------------------------------------------
doc.save(str(REPORT_OUT))
print(f"Saved {REPORT_OUT}")
