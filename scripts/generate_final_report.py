from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
import os


def setup_ieee_page(section):
    """Setup A4 page with IEEE-like margins."""
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(1.6)
    section.right_margin = Cm(1.6)
    section.top_margin = Cm(1.9)
    section.bottom_margin = Cm(2.0)


def set_columns(section, num=2, space_twips=720):
    """Set column layout for the section."""
    sect_pr = section._sectPr
    cols = sect_pr.xpath("./w:cols")
    if cols:
        col_el = cols[0]
    else:
        col_el = OxmlElement("w:cols")
        sect_pr.append(col_el)
    col_el.set(qn("w:num"), str(num))
    col_el.set(qn("w:space"), str(space_twips))


def set_normal_style(doc):
    """Set default document style."""
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(10)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(3)
    normal.paragraph_format.line_spacing = 1.0


def add_title(doc, text):
    """Add centered, non-bold title."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = False
    r.font.name = "Times New Roman"
    r.font.size = Pt(14)


def add_author_affiliations(doc):
    """Add author table."""
    table = doc.add_table(rows=5, cols=2)
    left = [
        "Milan Jovkic",
        "Departman za racunarske nauke",
        "Prirodno-matematicki fakultet",
        "Novi Sad, Srbija",
        "milan.jovkic@student.uns.ac.rs",
    ]
    right = [
        "Uros Petraskovic",
        "Departman za racunarske nauke",
        "Prirodno-matematicki fakultet",
        "Novi Sad, Srbija",
        "uros.petraskovic@student.uns.ac.rs",
    ]
    for i in range(5):
        c1 = table.cell(i, 0).paragraphs[0]
        c1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = c1.add_run(left[i])
        r1.font.name = "Times New Roman"
        r1.font.size = Pt(9)

        c2 = table.cell(i, 1).paragraphs[0]
        c2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = c2.add_run(right[i])
        r2.font.name = "Times New Roman"
        r2.font.size = Pt(9)


def add_section_heading(doc, text):
    """Add section heading (I., II., III., etc. - NOT bold)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    r.bold = False
    r.font.name = "Times New Roman"
    r.font.size = Pt(10)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)


def add_subsection_heading(doc, text):
    """Add sub-heading (A., B., C., etc.) - NOT bold."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    r.bold = False
    r.font.name = "Times New Roman"
    r.font.size = Pt(10)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)


def add_text(doc, text, align=WD_ALIGN_PARAGRAPH.JUSTIFY, italic=False):
    """Add body text."""
    p = doc.add_paragraph()
    p.alignment = align
    r = p.add_run(text)
    r.italic = italic
    r.font.name = "Times New Roman"
    r.font.size = Pt(10)
    return p


def add_caption(doc, text):
    """Add figure caption."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(9)


def add_figure(doc, fig_path, width_cm=8.0, caption_text=""):
    """Add figure with caption if file exists."""
    if os.path.exists(fig_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(fig_path, width=Cm(width_cm))
        if caption_text:
            add_caption(doc, caption_text)
    return p


def main():
    out_path = r"c:\Users\Win10\Documents\GitHub\Analiza-i-Obrada\Izvestaj\Natonal Football League player performance prediction using.docx"
    doc = Document()
    set_normal_style(doc)

    fig_base = r"c:\Users\Win10\Documents\GitHub\Analiza-i-Obrada\Izvestaj\figures"
    fig_counter = 1

    first = doc.sections[0]
    setup_ieee_page(first)
    set_columns(first, num=2)

    add_title(doc, "Predvidjanje performansi NFL igraca")
    add_author_affiliations(doc)

    add_section_heading(doc, "Apstrakt")
    add_text(
        doc,
        "U radu je prikazan kompletan eksperimentalni okvir za predvidjanje buducih performansi NFL igraca na cetiri "
        "kljucne ofanzivne pozicije: quarterback, running back, wide receiver i tight end. Motivacija je direktno povezana "
        "sa donosenjem odluka u uslovima ogranicenog budzeta i visoke konkurencije, gde pogresna procena igraca ima veliku "
        "finansijsku i sportsku cenu. Skup podataka formiran je kombinacijom sopstvenog procesa prikupljanja podataka sa "
        "platforme Pro Football Reference i javno dostupnog WR skupa sa HuggingFace izvora. Eksplorativna analiza ukazala je "
        "na strukturne nedostajuce vrednosti, poziciono razlicite raspodele cilja, izrazeniju multikolinearnost i asimetriju "
        "kod WR podataka. U skladu sa tim uvedene su ciljne promenljive po utakmici, lag atributi prve i druge istorijske "
        "sezone, enkodiranje nagrada i vremenski konzistentna evaluacija bez curenja podataka. Uporedjeni su linearni, "
        "regularizovani i ansambl modeli. Rezultati pokazuju da najbolji model zavisi od pozicije, pri cemu je TE pozicija "
        "najpredvidljivija, dok QB i RB zadrzavaju vecu neizvesnost usled jacih kontekstualnih uticaja. Dodatna walk forward "
        "kalibracija potvrdjuje vremensku robusnost pristupa i opravdava upotrebu predlozenog okvira kao osnove za naredna "
        "prosirenja ka multi output i sekvencijalnim modelima."
    )
    add_text(
        doc,
        "Kljucne reci: NFL analitika, predvidjanje performansi, masinsko ucenje, regresija, vremenska validacija",
        align=WD_ALIGN_PARAGRAPH.LEFT,
    )

    add_section_heading(doc, "I.   UVOD")
    add_text(
        doc,
        "Predvidjanje sportskog ucinka je problem sa izrazenom stohastickou komponenom, jer na rezultat uticu "
        "individualni kvalitet, timski sistem, uloga igraca, raspored protivnika i zdravstveno stanje. Kod NFL "
        "podataka dodatni izazov je sto tradicionalna sezonska agregacija cesto mesa kvalitet igre i dostupnost igraca "
        "na terenu. Zato je kljucno birati ciljnu promenljivu koja razdvaja ucinka od broja nastupa. U ovom radu "
        "koristimo metrike po utakmici, cimesmanjujemo pristrasnost prema igracima koji su propustili deo sezone."
    )
    add_text(
        doc,
        "Prakticna motivacija rada je dvoslojna. Prvi sloj je profesionalna primena u klubovima, gde preciznija procena "
        "buduceg ucinka pomaze pri odlukama o ugovorima, transferima i raspodeli budzeta. Drugi sloj je analiticka "
        "primena u fantasy okruzenju, gde su kvantitativni modeli vec standardni alat za procenu vrednosti igraca. "
        "U oba slucaja pozeljno je imati poziciono specificne modele, jer su obrasci performansi razliciti za QB, RB, WR i TE."
    )
    add_text(
        doc,
        "Glavni doprinos rada je integrisani i reproduktivan pipeline: od prikupljanja i ciscenja podataka, preko "
        "domenski vodenog inzenjeringa atributa, do rigorozne vremenske evaluacije. Posebno isticum odluku da se "
        "modeli treniraju odvojeno po pozicijama, uz iste principe validacije, kako bi poredjenje bilo metodoloski "
        "konzistentno. Ovaj pristup omogucava preciznije tumacenje gde su linearni obrasci dovoljni, a gde su potrebni "
        "nelinearni modeli. Struktura rada je sledeca: pregled literature, podaci i EDA nalazi, metodologija i evaluacioni "
        "protokol, rezultati/diskusija/analiza gresaka, ogranicenja i buduc rad, te zakljuci."
    )

    add_section_heading(doc, "II.   PREGLED LITERATURE")
    add_text(
        doc,
        "Savremeni radovi iz sportske analitike potvrdjuju da modeli masinskog ucenja nadmasuju jednostavne heuristike "
        "kada postoji dovoljno kvalitetnih atributa i pravilno definisan protokol validacije. Rad koji analizira "
        "predvidjanje procenta pobeda NFL timova pokazuje da ansambl i neuronski modeli ostvaruju boljuu "
        "tacnost od klasicnih formula, uz dodatnu interpretabilnost kroz SHAP analizu. Ta studija je vazna jer demonstrira "
        "da kvalitet ulaznih podataka i izbor evaluacije cesto imaju veci uticaj od same klase modela. Drugi relevantan smer "
        "daje rad o multi-output regresiji u vrhunskom sportu, gde se simultano predvidjabvise medjusobno povezanih performansnih metrika. "
        "Rezultat tog rada je da zajednicko modelovanje vise ciljeva moze poboljsati i preciznost i racunarsku efikasnost."
    )
    add_text(
        doc,
        "U domenu NFL i fantasy analitike, rad naglasava da poziciono razdvajanje nije samo prakticna izbor nego "
        "metodoloka nuznost. Metrike koje dominiraju kod QB nisu iste kao kod RB, niti kao kod WR/TE. Ovaj nalaz je direktno "
        "ugradjjen u nas dizajn. Na osnovu literature, nas pristup kombinuje tri principa: (1) vremenski konzistentnu validaciju, "
        "(2) domenski vodeno inzenjerstvo atributa i (3) uporednu analizu linearnih i nelinearnih modela. Time se postize "
        "ravnoteza izmedju prediktivne moci i interpretabilnosti."
    )

    add_section_heading(doc, "III.   PODACI I EKSPLORATIVNA ANALIZA")
    add_text(
        doc,
        "Podaci su prikupljeni iz dva izvora: Pro Football Reference i javni WR skup na HuggingFace platformi. "
        "Nakon harmonizacije, formirana su cetiri poziciona skupa: QB (845 x 188), RB (622 x 115), TE (351 x 70) i WR "
        "(5529 x 102). Ove dimenzije potvid uju da WR ima najveci broj zapisa, dok QB ima najsiri atributski prostor "
        "zbog kombinacije standardnih i naprednih passing metrika. Prvi nalaz EDA faze je da nedostajuce vrednosti nisu homogeni problem. "
        "Znacajan deo null vrednosti je strukturne prirode: pojedine metrike nisu istorijski postojale (npr. QBR pre 2006)."
    )

    add_figure(doc, os.path.join(fig_base, "target_dist_1.png"), width_cm=7.0,
               caption_text=f"Slika {fig_counter}. Distribucija ciljnih promenljivih po pozicijama")
    fig_counter += 1

    add_text(
        doc,
        "Pristup baziran na prostom globalnom imputiranju ocenjen je kao metodolski slab. Umesto toga, izvrsena je "
        "selekcija kolona i poziciono uslovljena tretman atributa, cime je smanjen sum i poboljsana stabilnost modela. "
        "Drugi nalaz odnosi se na raspodelu cilja. RB i TE pokazuju relativno stabilne i blize normalne distribucije, "
        "QB pokazuje bimodalnost (starteri naspram rezervnih igraca sa malim volumenom), dok WR ima izrazenu desnu "
        "asimetriju. Kod WR je zato primenjena log-transformacija cilja."
    )

    add_figure(doc, os.path.join(fig_base, "target_kde_1.png"), width_cm=7.0,
               caption_text=f"Slika {fig_counter}. KDE procena gustine ciljnih promenljivih")
    fig_counter += 1

    add_figure(doc, os.path.join(fig_base, "qb_bimodal_1.png"), width_cm=7.0,
               caption_text=f"Slika {fig_counter}. QB bimodalna distribucija: starteri vs rezervni igraci")
    fig_counter += 1

    add_text(
        doc,
        "Treci nalaz je izrazena multikolinearnost medu volume atributima. To je ocekivano: npr. vise targeta vodi ka vise "
        "hvatanja. Ovakav obrazac favorizuje regularizovane linearne modele i ansamble, dok nepenalizovani linearni modeli "
        "lako gube stabilnost koeficijenata. Pored numerickih atributa, analizirane su i ne-numericke kolone, posebno "
        "informacije o nagradama. Sirova tekstualna kolona nagrada transformisana je u pet binarnih atributa, cime su "
        "zadrzani informativni signali (Pro Bowl, All-Pro, MVP/OPoY)."
    )

    add_figure(doc, os.path.join(fig_base, "corr_qb_1.png"), width_cm=6.8,
               caption_text=f"Slika {fig_counter}a. Korelacijska matrica QB pozicije")
    fig_counter += 1

    add_figure(doc, os.path.join(fig_base, "corr_wr_1.png"), width_cm=6.8,
               caption_text=f"Slika {fig_counter}b. Korelacijska matrica WR pozicije")
    fig_counter += 1

    add_figure(doc, os.path.join(fig_base, "awards_pie_1.png"), width_cm=7.0,
               caption_text=f"Slika {fig_counter}. Ucestanost Pro Bowl statusa po pozicijama")
    fig_counter += 1

    add_section_heading(doc, "IV.   METODOLOGIJA")
    add_text(
        doc,
        "Ceo eksperiment dizajniran je tako da reprodukuje realan scenario predikcije buduce sezone. Prvo su formirane "
        "ciljne promenljive po utakmici za QB/RB/TE i odgovarajuca WR ciljna promenljiva sa log transformacijom. Nakon toga "
        "izvrseno je enkodiranje nagrada i selekcija kolona sa visokim udelom strukturnih null vrednosti. Za hvatanje vremenske "
        "dinamike uvedeni su lag atributi. Za svaku relevantnu numericku kolonu kreiraju se vrednosti iz prethodne sezone (lag1) "
        "i dve sezone unazad (lag2). Ovim se modelu daje istorijski kontekst bez curenja informacija iz sezone koja se predvidja."
    )

    add_figure(doc, os.path.join(fig_base, "wr_log_transform_1.png"), width_cm=7.0,
               caption_text=f"Slika {fig_counter}. WR log-transformacija: efekt na normalnost")
    fig_counter += 1

    add_text(
        doc,
        "Podela na trening i test izvedena je temporalno: sve sezone pre 2024 koriste se za obuku, sezona 2024 za test. "
        "Unutar treninga koriscen je TimeSeriesSplit sa pet foldova i GridSearchCV za izbor hiperparametara. Ovakav protokol "
        "je strozi od slucajnih podela i bolje meri generalizaciju kroz vreme. Modeli obuhvataju sirok spektar: Linear Regression, "
        "Ridge, Lasso, ElasticNet, k-Nearest Neighbors, Random Forest, XGBoost i LightGBM. Evaluacione metrike su MAE, RMSE i R2, "
        "jer zajedno pokrivaju prosecnu gresku, osetljivost na velika odstupanja i objasnjenu varijansu."
    )

    add_figure(doc, os.path.join(fig_base, "outliers_box_1.png"), width_cm=7.0,
               caption_text=f"Slika {fig_counter}. Outlier detektovanje: Z-score i IQR granice")
    fig_counter += 1

    ds_table = doc.add_table(rows=1, cols=4)
    ds_hdr = ds_table.rows[0].cells
    ds_hdr[0].text = "Pozicija"

        add_text(
            doc,
            "Redosled koraka je bitan: ako se imputacija i skaliranje urade pre vremenske podele, dobija se curenje informacija "
            "i nerealno optimistican rezultat. Kroz vremensku unakrsnu validaciju svi koraci se primenjuju samo na trening data "
            "za svaki fold, a zatim na odgovarajuci test skup. To osigurava da se model nikada ne vidi sa test periodom ni tokom "
            "kalibracije ni tokom selekcije hiperparametara. Dodatni nivo rigoroznosti postize se kroz walk-forward protokol: "
            "treniranje do 2020 i test 2021, zatim do 2021 i test 2022, do 2022 i test 2023, i do 2023 i test 2024. U svakom "
            "fold-u imputacija, skaliranje i optimizacija rade se iskljucivo na trening delu tog fold-a."
        )

        add_text(
            doc,
            "Transparentnost je unapredjena izborom metrika koje su medjusobno komplementarne. RMSE naglasava skupe greske i "
            "pokazuje osetljivost na ekstremne vrednosti, MAE daje robustniji pogled na tipican promasaj, a R2 omogucava "
            "relativno poredjenje izmedju modela unutar iste pozicije. Za WR je posebno vazno sto su metrike vracene na originalnu "
            "skalu nakon log transformacije, cime se izbegava pogresno tumacenje performansi modela transformisanom prostoru. "
            "Svi koraci su organizovani tako da mogu biti izvrseni kao jedinstven pipeline bez manuelnih intervencija."
        )
    ds_hdr[1].text = "Broj redova"
    ds_hdr[2].text = "Broj kolona"
    ds_hdr[3].text = "Ciljna promenljiva"
    for pos, rows, cols, target in [
        ("QB", "845", "188", "Passing Yds/Game"),
        ("RB", "622", "115", "Rushing Yds/Game"),
        ("WR", "5529", "102", "Rec Yds/Game (log)"),
        ("TE", "351", "70", "Rec Yds/Game"),
    ]:
        c = ds_table.add_row().cells
        c[0].text = pos
        c[1].text = rows
        c[2].text = cols
        c[3].text = target
    add_caption(doc, f"Tabela 1. Osnovna struktura podataka po pozicijama")

    add_section_heading(doc, "V.   REZULTATI I DISKUSIJA")
    add_text(
        doc,
        "Rezultati na test sezoni 2024 potvrdjuju da nema univerzalno najboljeg modela za sve pozicije. Za QB najbolji R2 "
        "postize XGBoost (0.145), za RB Random Forest (0.150), za TE Lasso (0.451), a za WR Random Forest (0.325). "
        "Najvazniiji zakljucak je da je TE pozicija najjednacinija i najpredvidljivija, dok QB i RB zadrzavaju nizu objasnjenu "
        "varijansu zbog jace uticaja faktora koji nisu u potpunosti sadrzani u tabelarnim atributima."
    )

    add_figure(doc, os.path.join(fig_base, "qb_scatter_1.png"), width_cm=7.0,
               caption_text=f"Slika {fig_counter}. QB starter analiza: games started i per-game metrike")
    fig_counter += 1

    add_figure(doc, os.path.join(fig_base, "qb_career_1.png"), width_cm=7.0,
               caption_text=f"Slika {fig_counter}. QB karijerna putanja: trend kroz godine")
    fig_counter += 1

    res_table = doc.add_table(rows=1, cols=5)
    rh = res_table.rows[0].cells
    rh[0].text = "Pozicija"
    rh[1].text = "Najbolji model"
    rh[2].text = "MAE"
    rh[3].text = "RMSE"
    rh[4].text = "R2"
    for row in [
        ("QB", "XGBoost", "30.84", "43.94", "0.145"),
        ("RB", "Random Forest", "19.44", "23.52", "0.150"),
        ("TE", "Lasso", "10.65", "11.59", "0.451"),
        ("WR", "Random Forest", "9.63", "12.61", "0.325"),
    ]:
        c = res_table.add_row().cells
        c[0].text = row[0]
        c[1].text = row[1]
        c[2].text = row[2]
        c[3].text = row[3]
        c[4].text = row[4]
    add_caption(doc, "Tabela 2. Performanse modela na test sezoni 2024")

    add_text(
        doc,
        "Kod QB rezultati pokazuju najvecu varijabilnost kroz sezone. Jedan deo greske dolazi iz nestabilnih promena u "
        "ofanzivnoj semi i dostupnosti kljucnih saigraca. Kod RB proizvodnja yardi zavisi od odnosa izmedju volumena nosenja "
        "i kvaliteta blokiranje. Cak i kada je igrac individualno stabilan, jedan deo varijanse dolazi iz timske strategije "
        "i situacionog pozivanja akcija. TE je pozicija sa najstabilnijim signalom jer se uloga igraca cesto razvija postepeno. "
        "WR ostvaruje umeren visok predvidljivosti zahvaljujuci velikom broju zapisa, ali ostaje osetljiv na asimeticne ekstreme."
    )

    add_figure(doc, os.path.join(fig_base, "rb_career_1.png"), width_cm=7.0,
               caption_text=f"Slika {fig_counter}. RB karijerna putanja")
    fig_counter += 1

    add_text(
        doc,
        "Walk-forward kalibracija dodatno potvrdjuje ove obrasce. Prosecno najbolji modeli su: QB XGBoost (AvgRMSE 51.564; "
        "AvgR2 0.176), RB ElasticNet (AvgRMSE 19.873; AvgR2 0.244), TE Ridge (AvgRMSE 13.052; AvgR2 0.502) i WR LightGBM "
        "(AvgRMSE 16.414; AvgR2 0.328). Analiza gresaka pokazuje tri dominantna obrasca: (1) sezone sa povredama gde modeli "
        "precenjuju igrace koji su istorijski stabilni; (2) transferi i promene sistema igre; (3) ekstremno produktivne sezone "
        "elitnih WR igraca. Prakticna implikacija je da modeli trebaju biti korisceni kao kvantitativni prior, ne kao jedini "
        "izvor odluke."

        add_text(
            doc,
            "Detaljnija analiza po podgrupama otkriva sledece: kod igraca sa naglim padom broja meceva, modeli najcesce precenjuju "
            "rezultat jer istorijski trend ne predvidja iznenadnu nedostupnost. Kod igraca koji su promenili tim greske su simetricenije, "
            "sto ukazuje da transfer moze da donese i skok i pad performansi. Kod vrhunskih WR-a sa ekstremno visokim volumenom, greske "
            "su vecinom potcenjivanja, sto je tipicno za sisteme sa izrazenim Pareto obrascem."
        )

        add_text(
            doc,
            "Operativna preporuka je da se za ove podgrupe koristi posebna post proces procena rizika. Osnovna predikcija ostaje "
            "centralna, ali u izvestaj treba ukljuciti indikator neizvesnosti koji je visi za transfer sezone, povratak posle povrede "
            "i ekstremne usage profile. Za prakticnu primenu u klubovima, predikcija po utakmici ima vecu analiticku vrednost od sezonskog "
            "zbira jer bolje povezuje predikciju sa taktickim odlukama i ugovornim procenama igraca."
        )
    )

    cal_table = doc.add_table(rows=1, cols=4)
    ch = cal_table.rows[0].cells
    ch[0].text = "Pozicija"
    ch[1].text = "Best Walk-Forward"
    ch[2].text = "AvgRMSE"
    ch[3].text = "AvgR2"
    for row in [
        ("QB", "XGBoost", "51.564", "0.176"),
        ("RB", "ElasticNet", "19.873", "0.244"),
        ("TE", "Ridge", "13.052", "0.502"),
        ("WR", "LightGBM", "16.414", "0.328"),
    ]:
        c = cal_table.add_row().cells
        c[0].text = row[0]
        c[1].text = row[1]
        c[2].text = row[2]
        c[3].text = row[3]
    add_caption(doc, "Tabela 3. Walk-forward kalibracija (2021-2024)")

    add_figure(doc, os.path.join(fig_base, "wr_log_outliers_1.png"), width_cm=7.0,
               caption_text=f"Slika {fig_counter}. WR outlier analiza: raw vs log vrednosti")
    fig_counter += 1

    add_subsection_heading(doc, "A.   Poziciono tumacenje")
    add_text(
        doc,
        "Detaljnije tumacenje po pozicijama otkriva zasto modeli imaju razlicitu granicu performansi. Kod QB, analiza starter "
        "vs. non-starter dinamike pokazuje kako se model obuka na mesovitoj populaciji bez eksplicitne separacije po ulozi. "
        "Promena trenerskog sistema i nivo zastite kvoterbeka mogu u jednoj sezoni promeniti i volumen i efikasnost. "
        "Istorijski atributi igraca jesu neophodni, ali cesto nisu dovoljni da potpuno objasne novu sezonsku dinamiku."
    )

    add_figure(doc, os.path.join(fig_base, "qb_trends_1.png"), width_cm=6.8,
               caption_text=f"Slika {fig_counter}a. QB trendovi kroz godine")
    fig_counter += 1

    add_figure(doc, os.path.join(fig_base, "qb_trends_2.png"), width_cm=6.8,
               caption_text=f"Slika {fig_counter}b. QB trendovi (nastavak)")
    fig_counter += 1

    add_figure(doc, os.path.join(fig_base, "rb_trends_1.png"), width_cm=6.8,
               caption_text=f"Slika {fig_counter}a. RB trendovi")
    fig_counter += 1

    add_figure(doc, os.path.join(fig_base, "wr_pareto_1.png"), width_cm=6.8,
               caption_text=f"Slika {fig_counter}b. WR Pareto analiza: 80/20 pravilo")
    fig_counter += 1

    add_subsection_heading(doc, "B.   Uticaj inzenjeringa atributa")
    add_text(
        doc,
        "Posebno znacajan doprinos performansi dolazi iz lag atributa. Bez vremenskih pomerenja model bi imao samo presecan "
        "pogled na sezonu. Lag1 i lag2 atributi omogucavaju modelima da nauče obrasce kontinuiteta i regresije ka proseku. "
        "Atribut Team_Changed uveden je da opise skokove i padove nakon transfera. Ovaj atribut poboljsava robusnost modela "
        "u tranzicionim sezonama."
    )

    add_figure(doc, os.path.join(fig_base, "qb_team_change_1.png"), width_cm=7.0,
               caption_text=f"Slika {fig_counter}. QB timske promene: efekt na performanse")
    fig_counter += 1

    add_subsection_heading(doc, "C.   Stabilnost i reproduktivnost")
    add_text(
        doc,
        "Reproduktivnost eksperimenta obezbedjenja je kroz strogo odvojene faze obrade: imputacija i skaliranje fituju se "
        "iskljucivo na trening skupu svakog fold-a, a zatim primenuju na test. Time je izbegnuto curenje informacija. "
        "Koriscenje vise metrika (MAE, RMSE, R2) omogucava uravnotzenu procenu modela. Walk-forward evaluacija pokazuje da "
        "pojedinacni test rezultat nikada ne treba tumaciti izolovano."
    )

    add_figure(doc, os.path.join(fig_base, "target_qq_1.png"), width_cm=7.0,
               caption_text=f"Slika {fig_counter}. Q-Q plot za normalnost proveru")
    fig_counter += 1

    add_section_heading(doc, "VI.   OGRANICENJA I BUDUC RAD")

        add_subsection_heading(doc, "D.   Poredjenje sa literaturom i prakticna vrednost")
        add_text(
            doc,
            "Poredjenje sa relevantnim radovima pokazuje dobru usaglasenost trendova. Kao i u literaturi koja analizira timske NFL "
            "ishode, ansambl modeli su konkurentni kada su ulazni atributi heterogeni i nelinearno povezani. Istovremeno, na "
            "pozicijama sa stabilnijim signalom regularizovani linearni modeli daju odlicne rezultate i bolju interpretabilnost. "
            "Ovaj nalaz je vazan jer potvrduje da izbor modela trebaju prilagoditi prirodi podataka, a ne pratiti univerzalni recept. "
            "U odnosu na radove koji se fokusiraju na fantasy poene, ovdje je prednost u tome sto se modeluju direktne sportske "
            "performanse po utakmici, sto ima vecu analiticku vrednost za klubove."
        )

        add_text(
            doc,
            "Za prakticnu primenu u klubovima, najvazniji ishod je mapa neizvesnosti po pozicijama, jer omogucava da se budzet "
            "raspodeli uz kontrolu rizika. Za skauting tim, korisna je identifikacija igraca kod kojih model vidi stabilan uzlazni trend "
            "i kod kojih istorijski signal ima visoku pouzdanost. Za trenerski stab, predikcija po utakmici je korisna jer bolje odrazava "
            "realan doprinos igraca u planiranju game plana. U fantasy analitici, model moze da sluzi kao objektivna osnova za rangiranje "
            "igraca, sa opreznošcu kod igraca sa ekstremnim profilom upotrebe."
        )
    add_text(
        doc,
        "Najvaznija ogranicenja su: (1) nedostajuci detaljni injury tokovi po nedeljama; (2) kvalitativni matchup podaci; "
        "(3) istorijska neujednacenost metrika. Buduc rad planiran je u cetiri pravca: (1) multi-output predikcija vise "
        "performansnih metrika; (2) sekvencijalni modeli sa punom vremenskom putanjom; (3) eksplicitno modelovanje rizika od "
        "povreda; (4) hijerarhijski modeli sa timskim i trenerskim kontekstom. Za prakticnu primenu u klubovima vazan je naredni "
        "korak operacionalizacija modela kroz periodicne projekcije sa verzionisanjem i standardizovanim izvestajima neizvesnosti."
    )

    add_section_heading(doc, "VII.   ZAKLJUCAK")
    add_text(
        doc,
        "Prikazani pristup potvrdjuje da je kvalitet predikcije u NFL domenu prvenstveno uslovljen pozicionim karakteristikama "
        "i ispravnim evaluacionim protokolom. Modeli koji ignorisu vremensku strukturu daju optimisticnu procenu i slabu generalizaciju. "
        "Suprotno tome, temporalna validacija i domenski vodjeni feature engineering omogucavaju realnije procene performansi u "
        "sledecoj sezoni. Rezultati jasno diferenciraju pozicije: TE je najstabilniji, WR umereno predvidljiv, QB i RB osetljiviji "
        "na spoljasnje faktore. Predlozeni okvir je dovoljno robustan za prakticnu upotrebu kao podrska odlucivanju, uz jasna "
        "ogranicenja i definisan plan unapredjenja."
    )

    add_section_heading(doc, "LITERATURA")
    references = [
        "[1] Autori, Advancing NFL win prediction, Sports Analytics Review, 2025.",
        "[2] Elimam et al., Multi-Output Regression in Sports, 2025.",
        "[3] Abadzic et al., Fantasy Football Prediction Analysis, 2024.",
        "[4] Pro Football Reference, https://www.pro-football-reference.com/",
        "[5] NFL WR Dataset, https://huggingface.co/datasets/SebastianAndreu/",
    ]
    for ref in references:
        add_text(doc, ref, align=WD_ALIGN_PARAGRAPH.LEFT)

    try:
        doc.save(out_path)
        saved_path = out_path
    except PermissionError:
        saved_path = out_path.replace(".docx", "_temp.docx")
        doc.save(saved_path)

    all_text = " ".join(p.text for p in doc.paragraphs if p.text.strip())
    word_count = len(all_text.split())
    print(f"✓ Saved: {saved_path}")
    print(f"✓ Words: {word_count}")
    print(f"✓ Figures: {fig_counter - 1}")


if __name__ == "__main__":
    main()
